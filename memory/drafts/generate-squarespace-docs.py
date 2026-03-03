#!/usr/bin/env python3
"""
Generates Squarespace-tailored resume + cover letter as Google Docs
via Drive API upload (converts .docx → Google Doc automatically).
"""

import json, io
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# ── Auth ───────────────────────────────────────────────────────────────────────
with open('/Users/jtsomwaru/.openclaw/workspace/config/google-oauth-token.json') as f:
    tok = json.load(f)

creds = Credentials(
    token=tok.get('token'),
    refresh_token=tok.get('refresh_token'),
    token_uri='https://oauth2.googleapis.com/token',
    client_id=tok.get('client_id'),
    client_secret=tok.get('client_secret'),
    scopes=tok.get('scopes', ['https://www.googleapis.com/auth/drive'])
)
drive = build('drive', 'v3', credentials=creds)

# ── Helpers ────────────────────────────────────────────────────────────────────
TEAL   = RGBColor(0x0d, 0x73, 0x77)
DARK   = RGBColor(0x1a, 0x1a, 0x2e)
GRAY   = RGBColor(0x55, 0x55, 0x55)

def set_font(run, size, bold=False, color=None, italic=False):
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def heading(doc, text, size=11, color=TEAL, bold=True, space_before=10, space_after=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text.upper())
    set_font(run, size, bold=bold, color=color)
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0d7377')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def job_title(doc, title, meta, space_before=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(title)
    set_font(r1, 10.5, bold=True, color=DARK)
    r2 = p.add_run(f'  ·  {meta}')
    set_font(r2, 9.5, italic=True, color=GRAY)

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.2)
    if bold_prefix:
        rb = p.add_run(bold_prefix + ' ')
        set_font(rb, 9.5, bold=True, color=DARK)
        r = p.add_run(text)
    else:
        r = p.add_run(text)
    set_font(r, 9.5, color=GRAY)

def body(doc, text, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    set_font(r, 10, color=GRAY)
    return p

def skill_row(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    rb = p.add_run(label + ':  ')
    set_font(rb, 9.5, bold=True, color=DARK)
    rv = p.add_run(value)
    set_font(rv, 9.5, color=GRAY)

def upload_to_drive(doc_obj, filename, folder_name):
    """Save docx to bytes, upload to Drive (convert to Google Doc), return URL."""
    buf = io.BytesIO()
    doc_obj.save(buf)
    buf.seek(0)

    # Find or create parent folder
    res = drive.files().list(
        q=f"name='{folder_name}' and mimeType='application/vnd.google-apps.folder' and trashed=false",
        fields='files(id)'
    ).execute()
    folders = res.get('files', [])

    # Find Eve — Drafts root
    root_res = drive.files().list(
        q="name='Eve — Drafts' and mimeType='application/vnd.google-apps.folder' and trashed=false",
        fields='files(id)'
    ).execute()
    root_folders = root_res.get('files', [])
    root_id = root_folders[0]['id'] if root_folders else None

    # Find Opticfy subfolder
    opticfy_res = drive.files().list(
        q=f"name='Opticfy' and mimeType='application/vnd.google-apps.folder' and '{root_id}' in parents and trashed=false",
        fields='files(id)'
    ).execute()
    opticfy_folders = opticfy_res.get('files', [])
    opticfy_id = opticfy_folders[0]['id'] if opticfy_folders else root_id

    meta = {
        'name': filename,
        'mimeType': 'application/vnd.google-apps.document',
        'parents': [opticfy_id] if opticfy_id else []
    }
    media = MediaIoBaseUpload(buf,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        resumable=False)
    f = drive.files().create(body=meta, media_body=media, fields='id,webViewLink').execute()
    return f.get('webViewLink')


# ══════════════════════════════════════════════════════════════════════════════
# RESUME
# ══════════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)

# ── Header ────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run('JT Somwaru')
set_font(r, 26, bold=True, color=DARK)

p2 = doc.add_paragraph()
p2.paragraph_format.space_after = Pt(3)
r2 = p2.add_run('AI Implementation Lead  ·  People AI & Automation Architect')
set_font(r2, 11.5, color=TEAL)

p3 = doc.add_paragraph()
p3.paragraph_format.space_after = Pt(2)
r3 = p3.add_run('New York, NY  ·  jsomwarux@yahoo.com  ·  linkedin.com/in/jtsomwaru  ·  github.com/jsomwarux')
set_font(r3, 9, color=GRAY)

# ── Summary ───────────────────────────────────────────────────────────────────
heading(doc, 'Professional Summary', space_before=8)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run(
    'Business Systems Analyst with 6 years at Spectrum Enterprise building enterprise system implementations '
    'through stakeholder observation, workflow analysis, and cross-functional coordination — the exact '
    'foundation this role requires. Now designing and deploying AI automation systems through Opticfy, '
    'with a track record of moving from POC to production across multiple AI-enabled tools and workflows. '
    'Deep hands-on experience with LLMs, AI agents, responsible AI practices, and vibe-coding functional '
    'prototypes. The person everyone asks about the latest AI — and who has shipped things that prove it.'
)
set_font(r, 9.8, color=GRAY)

# ── Core Competencies ─────────────────────────────────────────────────────────
heading(doc, 'Core Competencies')
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run(
    'AI-Directed Development (Vibe Coding)  ·  Responsible AI Practices  ·  POC → Production Delivery  ·  '
    'Workflow Observation & Analysis  ·  Multi-Agent Orchestration  ·  LLM Integration (Claude · GPT · Gemini · Grok)  ·  '
    'n8n Workflow Automation  ·  Business Systems Analysis  ·  Cross-Functional Stakeholder Alignment  ·  '
    'Prompt Engineering  ·  RAG Pipeline Design  ·  Playbook Creation  ·  Training & Coaching on AI Adoption'
)
set_font(r, 9.3, color=GRAY)

# ── Experience ────────────────────────────────────────────────────────────────
heading(doc, 'Professional Experience')

job_title(doc, 'Founder & AI Automation Consultant', 'Opticfy  ·  New York, NY  ·  2025 – Present')
bullet(doc, 'Defined the full AI consulting playbook — scripts, templates, evaluation frameworks, and implementation processes — '
       'from scratch; designed a 4-stage multi-agent delivery pipeline (Research → Workflow Builder → Presentation → Production).',
       bold_prefix='Playbook Creation:')
bullet(doc, 'Prototyped and deployed AI-enabled workflows using a "vibe-code first" methodology — '
       'rapidly building functional proof-of-concepts before committing to full production builds.',
       bold_prefix='Vibe Coding & POC:')
bullet(doc, 'Delivered first engagement: AI-powered construction progress tracking dashboard for Aya ($1,500) — '
       'automated field data collection, AI-generated status summaries, and stakeholder reporting.',
       bold_prefix='POC to Production:')
bullet(doc, 'Implemented responsible AI practices across all deployments: human oversight checkpoints, '
       'automated confidence scoring on AI decisions, immutable decision audit logs, and weekly '
       'AI explainability reporting summarizing all automated actions.',
       bold_prefix='Responsible AI:')
bullet(doc, 'Built and operate a 10-agent autonomous operations stack — market intelligence, job market scanning, '
       'portfolio analysis, and health tracking — demonstrating AI infrastructure thinking at scale.',
       bold_prefix='AI Infrastructure:')
bullet(doc, 'Coached clients and internal stakeholders on practical AI tool adoption — bridging the gap between '
       'AI capability and day-to-day workflow integration.',
       bold_prefix='Coaching:')

job_title(doc, 'Business Systems Analyst', 'Spectrum Enterprise / Charter Communications  ·  New York, NY  ·  2019 – 2025')
bullet(doc, 'Spent years doing the core of what this role asks: shadowing business teams, observing their workflows, '
       'and translating what they actually needed into implemented systems — for a Fortune 100 telecom enterprise.',
       bold_prefix='Workflow Observation:')
bullet(doc, 'Led product catalog configuration and system implementations across Salesforce and enterprise platforms — '
       'requirements gathering through deployment and ongoing iteration.',
       bold_prefix='System Implementation:')
bullet(doc, 'Translated complex business requirements from non-technical stakeholders into technical specifications '
       'for engineering and product teams; served as the primary business-to-technology bridge.',
       bold_prefix='Cross-Functional Bridge:')
bullet(doc, 'Coordinated multi-stakeholder projects across People, Tech, Operations, and Legal — managing competing '
       'priorities, driving alignment, and keeping implementations on track.',
       bold_prefix='Stakeholder Alignment:')
bullet(doc, 'Managed enterprise governance requirements: documentation, change control, compliance review, '
       'and audit trail maintenance for production system changes.',
       bold_prefix='Governance & Compliance:')

# ── AI Projects ───────────────────────────────────────────────────────────────
heading(doc, 'AI Products & Systems Built')

job_title(doc, 'People AI Operations Stack (OpenClaw)', '10 autonomous agents · 2025–Present', space_before=6)
bullet(doc, 'Designed and operate a fully autonomous AI operations stack: scheduling, memory management, '
       'agent coordination, cost tracking, and audit trail — always-on with zero manual restarts.')
bullet(doc, 'Includes responsible AI infrastructure: confidence scoring, human oversight routing, '
       'decision logging, and explainability reports — built from scratch.')

job_title(doc, 'Nash Satoshi — 4-LLM Cross-Validation Platform', '2025', space_before=6)
bullet(doc, '4-LLM ensemble (Claude + GPT + Gemini + Grok) that cross-validates AI outputs to produce '
       'consensus recommendations — demonstrates practical multi-model orchestration at production scale.')

job_title(doc, 'Mission Control — AI Operations Dashboard', '2025', space_before=6)
bullet(doc, 'Always-on Next.js + Convex dashboard for monitoring agents, tracking tasks (Kanban), '
       'viewing cost history, and maintaining a deployment audit trail — built via AI-directed development.')

job_title(doc, 'Salesforce Agentforce Deployments', '2025', space_before=6)
bullet(doc, 'Configured and deployed Employee Assistant (InternalCopilot) — 3 topics, 7 actions — '
       'and Ecommerce Service Agent (ExternalCopilot) using Agent Builder, Flows, and GenAiPlanner.')

# ── Skills ────────────────────────────────────────────────────────────────────
heading(doc, 'Technical Skills')
skill_row(doc, 'AI & Agents', 'Claude API, OpenAI GPT, Google Gemini, Grok, n8n (4-LLM ensemble patterns), '
          'multi-agent orchestration, RAG pipelines, prompt engineering (advanced)')
skill_row(doc, 'Development', 'Python (automation, API scripting, data processing), TypeScript/JavaScript (AI-directed), '
          'Git/GitHub, Next.js, SQLite')
skill_row(doc, 'Enterprise Systems', 'Salesforce (Agentforce, Flows, CRM data model), enterprise system configuration, '
          'requirements specification, process documentation')
skill_row(doc, 'Methodologies', 'AI-directed development (vibe coding), responsible AI practices, '
          'POC-to-production delivery, business process automation, stakeholder management')
skill_row(doc, 'Tools', 'Claude Code, Firecrawl, Brave Search API, Google Drive/Docs API, Convex, LaunchAgents')

# ── Upload Resume ─────────────────────────────────────────────────────────────
resume_url = upload_to_drive(doc, 'Resume — Squarespace People AI SA (Tailored)', 'Opticfy')
print(f'✅ Resume: {resume_url}')


# ══════════════════════════════════════════════════════════════════════════════
# COVER LETTER
# ══════════════════════════════════════════════════════════════════════════════
cl = Document()
for section in cl.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# Header
p = cl.add_paragraph()
r = p.add_run('JT Somwaru')
set_font(r, 13, bold=True, color=DARK)

p2 = cl.add_paragraph()
r2 = p2.add_run('jsomwarux@yahoo.com  ·  New York, NY  ·  linkedin.com/in/jtsomwaru')
set_font(r2, 9.5, color=GRAY)
p2.paragraph_format.space_after = Pt(18)

# Date + Addressee
p3 = cl.add_paragraph()
r3 = p3.add_run('February 24, 2026')
set_font(r3, 10, color=GRAY)
p3.paragraph_format.space_after = Pt(14)

p4 = cl.add_paragraph()
r4 = p4.add_run('Hiring Manager, People AI\nSquarespace\nNew York, NY')
set_font(r4, 10, color=GRAY)
p4.paragraph_format.space_after = Pt(14)

# Salutation
p5 = cl.add_paragraph()
r5 = p5.add_run('Dear Squarespace People Team,')
set_font(r5, 10.5, bold=True, color=DARK)
p5.paragraph_format.space_after = Pt(12)

# Para 1 — Hook
body(cl,
    'You used the phrase "vibe-code solutions" in a job description. That\'s the first time I\'ve seen '
    'my actual working methodology named as a job requirement. I\'ve been doing it for two years — '
    'directing Claude Code and specialized AI agents to prototype and ship functional tools from business '
    'requirements, without waiting for a roadmap or an engineering ticket. I\'m applying for the '
    'People AI Solutions Architect role because this is the job I\'ve been building toward.'
)

# Para 2 — BSA connection
body(cl,
    'I spent six years as a Business Systems Analyst at Spectrum Enterprise. The core of that job was '
    'exactly what your JD describes: shadowing teams, observing how they actually worked, and translating '
    'what they needed into implemented systems — not what they thought they needed, but what the workflow '
    'actually required. I did that across product operations, Salesforce implementations, and cross-functional '
    'projects involving People, Tech, Legal, and Operations stakeholders. That combination of workflow '
    'observation and cross-functional implementation is how I was trained to work, and it maps directly to '
    'what you\'re asking for.'
)

# Para 3 — AI work
body(cl,
    'Through Opticfy, I\'ve applied that foundation to AI. I\'ve built a 4-stage multi-agent delivery pipeline, '
    'shipped production AI tools using a POC-first approach, and deployed Salesforce Agentforce agents from '
    'Agent Builder configuration through full production. I built responsible AI infrastructure from scratch — '
    'human oversight checkpoints, automated confidence scoring, decision audit logs, and explainability reports '
    '— because I believe AI deployed in enterprise environments needs those guardrails, not as a checkbox but '
    'as part of the design. I also run a 10-agent autonomous operations stack that I built and maintain, '
    'which is less a portfolio piece and more proof that I actually live in this space.'
)

# Para 4 — People domain
body(cl,
    'People and HR is a newer domain for me specifically, and I\'ll say that directly. But the underlying '
    'pattern — shadow the COE, understand the real workflow, prototype the solution, iterate with users, '
    'move to production — is not new. It\'s what I\'ve done across every implementation I\'ve led. '
    'The Talent Acquisition and People Ops use cases are solvable with the same methodology I\'ve used '
    'to automate construction reporting, market intelligence, and business operations workflows. '
    'I expect a ramp on domain context, and I\'m comfortable getting it by doing what the JD says: '
    'spending time with the teams, observing how they work, and building what they actually need.'
)

# Para 5 — Close
body(cl,
    'I\'d welcome a conversation about the role. I can speak concretely to the POC-to-production work, '
    'the responsible AI practices I\'ve built, and what it looks like to write the playbook when one '
    'doesn\'t exist. Thank you for your time.'
)

# Sign off
p_sign = cl.add_paragraph()
p_sign.paragraph_format.space_before = Pt(16)
r_sign = p_sign.add_run('JT Somwaru')
set_font(r_sign, 10.5, bold=True, color=DARK)

# ── Upload Cover Letter ────────────────────────────────────────────────────────
cl_url = upload_to_drive(cl, 'Cover Letter — Squarespace People AI SA', 'Opticfy')
print(f'✅ Cover Letter: {cl_url}')
