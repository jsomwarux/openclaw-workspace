#!/usr/bin/env python3
"""
build_resume_docx.py — Generate properly formatted .docx resume and cover letter for JT.

Design spec:
  - Font: Calibri throughout
  - Name: 22pt bold, navy (#1B365D), centered
  - Contact line: 10pt dark gray (#555555), centered
  - Section headings: 12pt bold ALL CAPS navy, letter spacing, thin navy rule below
  - Job title: 11pt bold black, left-aligned
  - Company/dates: 10.5pt dark gray, right-aligned on same line
  - Bullets: 10.5pt black, 0.25" indent
  - Body text: 10.5pt black
  - Margins: 0.75" all sides
  - Colors: Navy #1B365D, Near-black #1A1A1A, Dark gray #555555
"""

import sys
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Color constants ────────────────────────────────────────────────────────────
NAVY     = RGBColor(0x1B, 0x36, 0x5D)
BLACK    = RGBColor(0x1A, 0x1A, 0x1A)
GRAY     = RGBColor(0x55, 0x55, 0x55)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
LT_GRAY  = RGBColor(0xF5, 0xF5, 0xF5)
BDR_GRAY = RGBColor(0xDD, 0xDD, 0xDD)


def set_run_color(run, color):
    run.font.color.rgb = color


def rgb_hex(color):
    """Convert RGBColor to hex string."""
    return f'{int(color[0]):02X}{int(color[1]):02X}{int(color[2]):02X}'

# Redefine colors as tuples for hex conversion
NAVY_HEX  = '1B365D'
BLACK_HEX = '1A1A1A'
GRAY_HEX  = '555555'
WHITE_HEX = 'FFFFFF'
LTGRAY_HEX = 'F5F5F5'

def add_horizontal_rule(para, color_hex=NAVY_HEX, width_pt=0.75):
    """Add a thin colored bottom border to a paragraph (acts as a divider rule)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(int(width_pt * 8)))  # sz in eighths of a point
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_para_spacing(para, before_pt=0, after_pt=0, line_spacing=None):
    pPr = para._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), str(int(before_pt * 20)))
    spacing.set(qn('w:after'), str(int(after_pt * 20)))
    if line_spacing:
        spacing.set(qn('w:line'), str(int(line_spacing * 240)))
        spacing.set(qn('w:lineRule'), 'auto')
    pPr.append(spacing)


def set_tab_stop(para, position_inches, alignment='right'):
    """Add a right tab stop for right-aligned dates on same line as left text."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), alignment)
    tab.set(qn('w:pos'), str(int(position_inches * 1440)))  # twips
    tabs.append(tab)
    pPr.append(tabs)


def set_indent(para, left_inches=0, hanging_inches=0):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(int(left_inches * 1440)))
    if hanging_inches:
        ind.set(qn('w:hanging'), str(int(hanging_inches * 1440)))
    pPr.append(ind)


def add_section_heading(doc, text):
    """Section heading: 12pt bold ALL CAPS navy + thin navy rule below."""
    para = doc.add_paragraph()
    set_para_spacing(para, before_pt=10, after_pt=3)
    run = para.add_run(text.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = NAVY
    # Letter spacing via rPr/w:spacing
    rPr = run._r.get_or_add_rPr()
    spacing_el = OxmlElement('w:spacing')
    spacing_el.set(qn('w:val'), '20')  # 1pt = 20 twips
    rPr.append(spacing_el)
    add_horizontal_rule(para)
    return para


def add_job_header(doc, title, location, company, dates, content_width=7.0):
    """Job title + location on line 1, company + dates on line 2."""
    # Line 1: Title (bold black left) + Location (gray right)
    p1 = doc.add_paragraph()
    set_para_spacing(p1, before_pt=8, after_pt=1)
    set_tab_stop(p1, content_width)
    r1 = p1.add_run(title)
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = BLACK
    if location:
        tab_run = p1.add_run('\t')
        loc_run = p1.add_run(location)
        loc_run.font.name = 'Calibri'
        loc_run.font.size = Pt(10.5)
        loc_run.font.color.rgb = GRAY

    # Line 2: Company (gray left) + Dates (gray right)
    p2 = doc.add_paragraph()
    set_para_spacing(p2, before_pt=0, after_pt=1)
    set_tab_stop(p2, content_width)
    comp_run = p2.add_run(company)
    comp_run.font.name = 'Calibri'
    comp_run.font.size = Pt(10.5)
    comp_run.font.color.rgb = GRAY
    if dates:
        tab_run = p2.add_run('\t')
        dates_run = p2.add_run(dates)
        dates_run.font.name = 'Calibri'
        dates_run.font.size = Pt(10.5)
        dates_run.font.color.rgb = GRAY


def add_bullet(doc, text):
    """Standard bullet point."""
    para = doc.add_paragraph(style='List Bullet')
    set_para_spacing(para, before_pt=0, after_pt=2)
    # Clear default bullet formatting and apply ours
    para.clear()
    # Use actual bullet character
    run = para.add_run('• ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    # Remove list style to use our custom bullet
    para.style = doc.styles['Normal']
    set_indent(para, left_inches=0.25)
    set_para_spacing(para, before_pt=0, after_pt=2)
    return para


def add_skills_line(doc, category, skills_text):
    """Skills: 'Category: skill1, skill2' with bold category."""
    para = doc.add_paragraph()
    set_para_spacing(para, before_pt=0, after_pt=3)
    cat_run = para.add_run(category + ': ')
    cat_run.font.name = 'Calibri'
    cat_run.font.size = Pt(10.5)
    cat_run.font.bold = True
    cat_run.font.color.rgb = BLACK
    skills_run = para.add_run(skills_text)
    skills_run.font.name = 'Calibri'
    skills_run.font.size = Pt(10.5)
    skills_run.font.bold = False
    skills_run.font.color.rgb = BLACK
    return para


def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)


def add_builds_table(doc, builds):
    """
    DEPRECATED — DO NOT USE. ATS systems cannot parse tables and will drop all content.
    Selected Builds must use inline text format (add_builds_inline). This function is
    kept only as a reference for the visual spec; it is never called.
    builds: list of (project, description, status)
    """
    raise RuntimeError("add_builds_table is deprecated — use inline text format for ATS compatibility.")
    table = doc.add_table(rows=1 + len(builds), cols=3)
    table.style = 'Table Grid'

    # Remove all borders and set custom ones
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement('w:tcBdr')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                bdr = OxmlElement(f'w:{side}')
                bdr.set(qn('w:val'), 'nil')
                tcBdr.append(bdr)
            tcPr.append(tcBdr)

    # Column widths (approx): 1.8" | 4.0" | 1.2"
    col_widths = [Inches(1.8), Inches(4.0), Inches(1.2)]
    for i, col in enumerate(table.columns):
        for cell in col.cells:
            cell.width = col_widths[i]

    # Header row
    header_row = table.rows[0]
    header_data = ['Project', 'Description', 'Status']
    for i, cell in enumerate(header_row.cells):
        set_cell_bg(cell, '1B365D')
        para = cell.paragraphs[0]
        para.clear()
        set_para_spacing(para, before_pt=4, after_pt=4)
        run = para.add_run(header_data[i])
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = WHITE

    # Data rows
    for row_idx, (project, description, status) in enumerate(builds):
        row = table.rows[row_idx + 1]
        bg = 'F5F5F5' if row_idx % 2 == 0 else 'FFFFFF'
        data = [project, description, status]
        for i, cell in enumerate(row.cells):
            set_cell_bg(cell, bg)
            para = cell.paragraphs[0]
            para.clear()
            set_para_spacing(para, before_pt=4, after_pt=4)
            run = para.add_run(data[i])
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = BLACK


def parse_resume_md(md_path):
    """
    Parse a structured resume .md file into sections.
    Expected headings: PROFESSIONAL SUMMARY, SKILLS, EXPERIENCE, KEY PROJECTS, EDUCATION, TECHNICAL STACK
    Returns a dict with parsed content.
    """
    import re

    with open(md_path) as f:
        raw = f.read()

    lines = raw.splitlines()
    data = {
        'name': '',
        'contact': '',
        'summary': '',
        'skills': [],       # list of (category, text)
        'experience': [],   # list of {title, company, location, dates, blurb, bullets}
        'projects': [],     # list of (name, description)
        'education': [],    # list of lines
        'tech_stack': [],   # list of (category, text)
    }

    # Name = first H1
    for line in lines:
        if line.startswith('# '):
            data['name'] = line[2:].strip()
            break

    # Contact = second non-empty line after name
    found_name = False
    for line in lines:
        if line.startswith('# '):
            found_name = True
            continue
        if found_name and line.strip():
            data['contact'] = line.strip()
            break

    # Parse sections by H2 headings
    current_section = None
    current_job = None
    buffer = []

    def flush_job():
        if current_job:
            current_job['bullets'] = [b for b in buffer]
            data['experience'].append(current_job)

    for i, line in enumerate(lines):
        # Section headings
        if line.startswith('## '):
            if current_job:
                flush_job()
                current_job = None
                buffer = []
            current_section = line[3:].strip().upper()
            buffer = []
            continue

        # Job headings (H3)
        if line.startswith('### ') and current_section == 'EXPERIENCE':
            if current_job:
                flush_job()
                buffer = []
            title = line[4:].strip()
            # Next non-empty lines: **Company**, location, dates
            current_job = {'title': title, 'company': '', 'location': '', 'dates': '', 'blurb': '', 'bullets': []}
            buffer = []
            continue

        if current_section == 'PROFESSIONAL SUMMARY':
            if line.strip() and not line.startswith('---'):
                data['summary'] += (' ' if data['summary'] else '') + line.strip()

        elif current_section == 'SKILLS':
            if line.strip() and not line.startswith('---'):
                # Format: **Category:** text
                m = re.match(r'\*\*(.+?):\*\*\s*(.*)', line.strip())
                if m:
                    data['skills'].append((m.group(1), m.group(2)))

        elif current_section == 'EXPERIENCE' and current_job is not None:
            stripped = line.strip()
            if not stripped or stripped == '---':
                continue
            # Company line: **Company**, Location, Dates
            if stripped.startswith('**') and not current_job['company']:
                m = re.match(r'\*\*(.+?)\*\*,?\s*(.*)', stripped)
                if m:
                    rest = m.group(2)
                    parts = [p.strip() for p in rest.split(',')]
                    current_job['company'] = m.group(1)
                    if len(parts) >= 2:
                        current_job['location'] = parts[0]
                        current_job['dates'] = parts[-1]
                    elif len(parts) == 1:
                        current_job['dates'] = parts[0]
            # Blurb (italic or plain paragraph before bullets)
            elif not stripped.startswith('-') and not current_job.get('blurb') and current_job['company']:
                current_job['blurb'] = stripped
            # Bullets
            elif stripped.startswith('- '):
                buffer.append(stripped[2:])

        elif current_section == 'KEY PROJECTS':
            if line.strip() and not line.startswith('---'):
                # Format: **Name** (status): description
                m = re.match(r'\*\*(.+?)\*\*\s*(\([^)]*\))?:?\s*(.*)', line.strip())
                if m:
                    name = m.group(1)
                    desc = m.group(3)
                    data['projects'].append((name, desc))

        elif current_section == 'EDUCATION':
            if line.strip() and not line.startswith('---'):
                data['education'].append(line.strip())

        elif current_section == 'TECHNICAL STACK':
            if line.strip() and not line.startswith('---'):
                m = re.match(r'\*\*(.+?):\*\*\s*(.*)', line.strip())
                if m:
                    data['tech_stack'].append((m.group(1), m.group(2)))

    # Flush last job
    if current_job:
        flush_job()

    return data


def parse_cover_letter_md(md_path):
    """
    Parse a structured cover letter .md file.
    Extracts: recipient company name, body paragraphs.
    """
    with open(md_path) as f:
        raw = f.read()

    lines = raw.splitlines()
    data = {'company': 'Writer', 'recipient_line': 'Hiring Team', 'paragraphs': []}

    # Try to extract company from H1 or first metadata line
    for line in lines:
        if 'Applying for:' in line or 'applying for:' in line.lower():
            import re
            m = re.search(r'\|\s*(.+)$', line)
            if m:
                data['company'] = m.group(1).strip()
            break

    # Extract recipient block (lines after --- that have "Hiring" or company)
    for line in lines:
        if line.strip().startswith('Hiring'):
            data['recipient_line'] = line.strip()
            break

    # Extract body paragraphs: non-empty lines after second --- that are actual prose
    # (not headers, not metadata, not closing)
    separators = [i for i, l in enumerate(lines) if l.strip() == '---']
    body_start = separators[1] + 1 if len(separators) >= 2 else 0

    body_lines = lines[body_start:]
    paragraphs = []
    current = []

    for line in body_lines:
        stripped = line.strip()
        if not stripped:
            if current:
                paragraphs.append(' '.join(current))
                current = []
        elif stripped == '---':
            if current:
                paragraphs.append(' '.join(current))
                current = []
            break
        elif stripped.startswith('Jon Trevor') or stripped.startswith('jtsomwaru'):
            # Closing block — stop
            if current:
                paragraphs.append(' '.join(current))
            break
        else:
            current.append(stripped)

    if current:
        paragraphs.append(' '.join(current))

    # Filter out blank or very short lines (metadata artifacts)
    data['paragraphs'] = [p for p in paragraphs if len(p) > 20]

    return data


def build_resume(output_path, resume_md=None):
    # ── Load data from markdown if provided (preferred) ───────────────────────
    md = None
    if resume_md:
        md = parse_resume_md(resume_md)

    doc = Document()

    # ── Page setup ────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(0.75)
    section.right_margin  = Inches(0.75)
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    content_width = 7.0  # 8.5 - 0.75 - 0.75

    # ── Default paragraph style ───────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.font.color.rgb = BLACK

    # ── HEADER ───────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(name_para, before_pt=0, after_pt=4)
    name_run = name_para.add_run(md['name'] if md else 'Jon Trevor Somwaru')
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = NAVY

    # Contact line — standardized regardless of md (always includes jtsomwaru.com)
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para_spacing(contact_para, before_pt=0, after_pt=6)
    contact_run = contact_para.add_run(
        'New York City · jtsomwaru@gmail.com · '
        'linkedin.com/in/jon-trevor-somwaru · jtsomwaru.com'
    )
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = GRAY
    add_horizontal_rule(contact_para)

    # ── SUMMARY ──────────────────────────────────────────────────────────────
    add_section_heading(doc, 'Professional Summary')
    summary_para = doc.add_paragraph()
    set_para_spacing(summary_para, before_pt=4, after_pt=6, line_spacing=1.15)
    summary_text = md['summary'] if md else (
        'AI Implementation Consultant and Solutions Architect with 6 years of enterprise delivery experience. '
        'Specializes in generative AI architecture, multi-agent orchestration, and production AI deployment.'
    )
    summary_run = summary_para.add_run(summary_text)
    summary_run.font.name = 'Calibri'
    summary_run.font.size = Pt(10.5)
    summary_run.font.color.rgb = BLACK

    # ── EXPERIENCE ────────────────────────────────────────────────────────────
    add_section_heading(doc, 'Experience')

    if md and md['experience']:
        for job in md['experience']:
            add_job_header(doc,
                title=job['title'],
                location=job.get('location', 'New York City'),
                company=job['company'],
                dates=job['dates'],
                content_width=content_width
            )
            for b in job['bullets']:
                add_bullet(doc, b)
    else:
        # Fallback hardcoded experience (kept as safety net)
        add_job_header(doc,
            title='AI Implementation Consultant',
            location='New York City',
            company='Independent Consulting',
            dates='2025–Present',
            content_width=content_width
        )
        for b in [
            'Architected and deployed AgentGuard, a confidence-gated AI governance layer with automated routing at a 70% confidence threshold and a full EEOC-compliant audit trail. Live at agentguard-delta.vercel.app.',
            'Built and operate 35 autonomous cron jobs in production covering prospect research, outreach pipelines, content generation, market intelligence, and cost monitoring.',
            'Delivered a $1,500 operational intelligence dashboard for Aya (NYC construction firm); client commissioned $3,500 in follow-on projects immediately after delivery.',
        ]:
            add_bullet(doc, b)

    # ── SKILLS (from md) — rendered AFTER experience ─────────────────────────
    if md and md['skills']:
        add_section_heading(doc, 'Skills')
        for cat, text in md['skills']:
            add_skills_line(doc, cat, text)

    # ── KEY PROJECTS ──────────────────────────────────────────────────────────
    projects_to_render = md['projects'] if md and md['projects'] else [
        ('AgentGuard', 'Confidence-gated AI governance layer. Live at agentguard-delta.vercel.app.'),
        ('Nash Satoshi', 'Crypto intelligence platform built on a 4-LLM ensemble architecture.'),
    ]
    if projects_to_render:
        add_section_heading(doc, 'Key Projects')
        for name, desc in projects_to_render:
            para = doc.add_paragraph()
            set_para_spacing(para, before_pt=3, after_pt=3, line_spacing=1.1)
            # Strip trailing colon from name if markdown already has it
            clean_name = name.rstrip(':').rstrip()
            name_run = para.add_run(clean_name + ': ')
            name_run.font.name = 'Calibri'
            name_run.font.size = Pt(10.5)
            name_run.font.bold = True
            name_run.font.color.rgb = BLACK
            desc_run = para.add_run(desc)
            desc_run.font.name = 'Calibri'
            desc_run.font.size = Pt(10.5)
            desc_run.font.bold = False
            desc_run.font.color.rgb = BLACK

    # ── TECHNICAL STACK ───────────────────────────────────────────────────────
    if md and md['tech_stack']:
        add_section_heading(doc, 'Technical Stack')
        for cat, text in md['tech_stack']:
            add_skills_line(doc, cat, text)

    # ── EDUCATION ─────────────────────────────────────────────────────────────
    add_section_heading(doc, 'Education')
    edu_lines = md['education'] if md and md['education'] else [
        'Bachelor of Science, Sport Management | Ithaca College | 2018',
        'Data Analytics Certificate | Northeastern University | 2019'
    ]
    import re as _re
    for edu_line in edu_lines:
        edu_para = doc.add_paragraph()
        set_para_spacing(edu_para, before_pt=4, after_pt=2)
        line = edu_line.strip()
        # Format 1: **School**, degree, year
        bold_m = _re.match(r'\*\*(.+?)\*\*,?\s*(.*)', line)
        # Format 2: Degree | School | Year  (pipe-separated, no markdown bold)
        pipe_m = _re.match(r'(.+?)\s*\|\s*(.+?)\s*\|\s*(.*)', line) if not bold_m else None
        if bold_m:
            r1 = edu_para.add_run(bold_m.group(1))
            r1.font.name = 'Calibri'; r1.font.size = Pt(10.5)
            r1.font.bold = True; r1.font.color.rgb = BLACK
            r2 = edu_para.add_run(', ' + bold_m.group(2) if bold_m.group(2) else '')
            r2.font.name = 'Calibri'; r2.font.size = Pt(10.5)
            r2.font.bold = False; r2.font.color.rgb = GRAY
        elif pipe_m:
            # "Degree | School | Year" → bold school, normal rest
            degree, school, year = pipe_m.group(1).strip(), pipe_m.group(2).strip(), pipe_m.group(3).strip()
            r1 = edu_para.add_run(school)
            r1.font.name = 'Calibri'; r1.font.size = Pt(10.5)
            r1.font.bold = True; r1.font.color.rgb = BLACK
            rest = f', {degree}' + (f', {year}' if year else '')
            r2 = edu_para.add_run(rest)
            r2.font.name = 'Calibri'; r2.font.size = Pt(10.5)
            r2.font.bold = False; r2.font.color.rgb = GRAY
        else:
            r = edu_para.add_run(line)
            r.font.name = 'Calibri'; r.font.size = Pt(10.5)
            r.font.color.rgb = BLACK

    doc.save(output_path)
    print(f'✅ Resume saved: {output_path}')


def build_cover_letter(output_path, cover_letter_md=None):
    # ── Load data from markdown if provided ───────────────────────────────────
    cl = None
    if cover_letter_md:
        cl = parse_cover_letter_md(cover_letter_md)

    doc = Document()

    section = doc.sections[0]
    section.page_width  = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = BLACK

    # ── HEADER ───────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(name_para, before_pt=0, after_pt=2)
    name_run = name_para.add_run('Jon Trevor Somwaru')
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(20)
    name_run.font.bold = True
    name_run.font.color.rgb = NAVY

    contact_para = doc.add_paragraph()
    set_para_spacing(contact_para, before_pt=0, after_pt=12)
    contact_run = contact_para.add_run(
        'jtsomwaru@gmail.com · jtsomwaru.com · linkedin.com/in/jon-trevor-somwaru'
    )
    contact_run.font.name = 'Calibri'
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = GRAY
    add_horizontal_rule(contact_para)

    # Date + recipient block
    date_para = doc.add_paragraph()
    set_para_spacing(date_para, before_pt=8, after_pt=2)
    date_run = date_para.add_run(datetime.now().strftime('%B %Y'))
    date_run.font.name = 'Calibri'
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = GRAY

    recip_para = doc.add_paragraph()
    set_para_spacing(recip_para, before_pt=0, after_pt=12)
    recip_line = (cl['recipient_line'] if cl else 'Hiring Team')
    recip_run = recip_para.add_run(recip_line)
    recip_run.font.name = 'Calibri'
    recip_run.font.size = Pt(11)
    recip_run.font.color.rgb = GRAY

    # ── BODY PARAGRAPHS ──────────────────────────────────────────────────────
    # Use parsed paragraphs from markdown when available — this is the source of truth
    body_paras = cl['paragraphs'] if cl and cl['paragraphs'] else [
        'Cover letter content not found. Please provide a --cover-letter-md file.'
    ]

    for body_text in body_paras:
        p = doc.add_paragraph()
        set_para_spacing(p, before_pt=0, after_pt=10, line_spacing=1.15)
        run = p.add_run(body_text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = BLACK

    # ── CLOSING ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    set_para_spacing(p, before_pt=6, after_pt=2)
    run = p.add_run('Sincerely,')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = BLACK

    p2 = doc.add_paragraph()
    set_para_spacing(p2, before_pt=14, after_pt=1)
    r1 = p2.add_run('Jon Trevor Somwaru')
    r1.font.name = 'Calibri'
    r1.font.size = Pt(11)
    r1.font.bold = True
    r1.font.color.rgb = BLACK

    p3 = doc.add_paragraph()
    set_para_spacing(p3, before_pt=0, after_pt=0)
    r2 = p3.add_run('jtsomwaru@gmail.com · jtsomwaru.com')
    r2.font.name = 'Calibri'
    r2.font.size = Pt(10)
    r2.font.color.rgb = GRAY

    doc.save(output_path)
    print(f'✅ Cover letter saved: {output_path}')


if __name__ == '__main__':
    import argparse
    parser = argparse.ArgumentParser(description='Build formatted .docx resume and cover letter from markdown source files.')
    parser.add_argument('--type', choices=['resume', 'cover_letter', 'both'], default='both')
    parser.add_argument('--output-dir', default='/tmp')
    parser.add_argument('--resume-md', default=None,
                        help='Path to resume markdown file (source of truth for content). '
                             'If omitted, falls back to hardcoded content.')
    parser.add_argument('--cover-letter-md', default=None,
                        help='Path to cover letter markdown file (source of truth for content). '
                             'If omitted, falls back to placeholder text.')
    args = parser.parse_args()

    if args.type in ('resume', 'both'):
        build_resume(f'{args.output_dir}/jt-somwaru-resume.docx', resume_md=args.resume_md)
    if args.type in ('cover_letter', 'both'):
        build_cover_letter(f'{args.output_dir}/jt-somwaru-cover-letter.docx', cover_letter_md=args.cover_letter_md)
